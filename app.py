# app.py (Versio 7.2 - Vakaa jäsennys ja parannettu relevanssi)
import re
from collections import defaultdict
from io import BytesIO
import docx
import streamlit as st

from logic import etsi_merkityksen_mukaan


# --- Sivun asetukset ---
st.set_page_config(
    page_title="Raamattu-tutkija v4",
    page_icon="📚",
    layout="wide"
)


# --- Apufunktiot ---
def lue_syote_data(syote_data):
    """
    Jäsentää syötetiedon älykkäästi, yhdistäen otsikot ja selitykset
    parhaan hakutuloksen saavuttamiseksi.
    """
    if not syote_data:
        return None, None, None, None

    if isinstance(syote_data, str):
        sisalto = syote_data
    else:
        sisalto = syote_data.getvalue().decode("utf-8")
    sisalto = sisalto.replace('\r\n', '\n')

    paaotsikko_m = re.search(r"^(.*?)\n", sisalto)
    paaotsikko = paaotsikko_m.group(1).strip() if paaotsikko_m else ""
    hakulauseet, otsikot, sl_teksti = {}, {}, ""

    # STRATEGIA 1: Etsi uudempi sisällysluettelomalli
    sl_m = re.search(
        r"sisällysluettelo:\s*\n(.*?)(?=\n\n^\s*\d\.|\Z)",
        sisalto, re.I | re.S
    )
    if sl_m:
        sl_teksti = sl_m.group(1).strip()
        # Jaetaan sisällysluettelo osiin numeroidun otsikon perusteella
        osio_rivit = re.split(r'\n(?=\s*\d[\d\.]*\s)', sl_teksti)
        for rivi in osio_rivit:
            rivi = rivi.strip()
            if not rivi:
                continue
            match = re.match(r"^\s*(\d[\d\.]*)\s*(.*)", rivi, re.S)
            if match:
                osio_nro = match.group(1).strip().rstrip('.')
                koko_teksti = match.group(2).strip()
                otsikot[osio_nro] = koko_teksti.split('\n')[0]
                hakulauseet[osio_nro] = koko_teksti.replace('\n', ' ')
        return paaotsikko, sl_teksti, hakulauseet, otsikot

    # STRATEGIA 2: Etsi vanha Teema-malli
    vanha_sl_m = re.search(
        r"SISÄLLYSLUETTELO JA HAKUSANAT\n(.*?)(?=\n\n^\d\.|\Z)", sisalto, re.S
    )
    if vanha_sl_m:
        sl_teksti = vanha_sl_m.group(1).strip()

    p = re.compile(
        r"^(\d[\d\.]*)\s*.*?\n\nTeema:\s*(.*?)(?=\n\n^\d|\Z)", re.S | re.M
    )
    for osio_nro, teema in p.findall(sisalto):
        osio_nro = osio_nro.strip().rstrip('.')
        teema = teema.strip().replace('\n', ' ')
        hakulauseet[osio_nro] = teema
        otsikko_m = re.search(
            r"^{}\.?\s*(.*)".format(re.escape(osio_nro)), sl_teksti, re.M
        )
        otsikot[osio_nro] = (
            otsikko_m.group(1).strip() if otsikko_m else f"Osio {osio_nro}"
        )
    return paaotsikko, sl_teksti, hakulauseet, otsikot


def luo_raportti_md(sl, jae_kartta):
    """Muotoilee kerätyt jakeet siistiksi Markdown-raportiksi."""
    raportti_osat = ["# Raamattu-tutkijan Raportti"]
    if sl:
        raportti_osat.append(f"## Sisällysluettelo\n\n```\n{sl}\n```\n\n---")

    # KORJATTU JÄRJESTYS: Vankka logiikka, joka estää kaatumisen
    sorted_osiot = sorted(
        jae_kartta.items(),
        key=lambda item: [int(p) for p in item[0].split('.') if p.strip().isdigit()]
    )
    for osio_nro, data in sorted_osiot:
        otsikko = data.get("otsikko", f"Osio {osio_nro}")
        taso = osio_nro.count('.') + 2
        raportti_osat.append(f"\n{'#' * taso} {osio_nro} {otsikko}\n")
        jakeet = data.get("jakeet", [])
        if not jakeet:
            raportti_osat.append("*Ei löytynyt jakeita tähän osioon.*\n")
        else:
            for jae in jakeet:
                raportti_osat.append(f"- **{jae['viite']}**: {jae['teksti']}")
            raportti_osat.append("")
    return "\n".join(raportti_osat)


def luo_raportti_doc(sl, jae_kartta):
    """Luo raportista .docx-version muistiin."""
    doc = docx.Document()
    doc.add_heading("Raamattu-tutkijan Raportti", level=0)
    if sl:
        doc.add_heading("Sisällysluettelo", level=1)
        doc.add_paragraph(sl)
    # KORJATTU JÄRJESTYS: Vankka logiikka, joka estää kaatumisen
    sorted_osiot = sorted(
        jae_kartta.items(),
        key=lambda item: [int(p) for p in item[0].split('.') if p.strip().isdigit()]
    )
    for osio_nro, data in sorted_osiot:
        otsikko = data.get("otsikko", f"Osio {osio_nro}")
        taso = min(osio_nro.count('.') + 1, 4)
        doc.add_heading(f"{osio_nro} {otsikko}", level=taso)
        jakeet = data.get("jakeet", [])
        if not jakeet:
            doc.add_paragraph("Ei löytynyt jakeita tähän osioon.")
        else:
            for jae in jakeet:
                p = doc.add_paragraph()
                p.add_run(f"{jae['viite']}: ").bold = True
                p.add_run(jae['teksti'])
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


def main():
    """Sovelluksen pääohjelma ja käyttöliittymän rakentaminen."""
    st.title("📚 Raamattu-tutkija v4.0 - Merkityshaku")

    for key in ['processing_complete', 'final_report_md', 'final_report_doc']:
        if key not in st.session_state:
            st.session_state[key] = (
                False if key == 'processing_complete' else ""
            )

    input_method = st.radio(
        "Valitse syöttötapa:",
        ("Lataa tiedosto", "Kirjoita tai liitä teksti"),
        horizontal=True
    )
    syote_data = None
    if input_method == "Lataa tiedosto":
        uploaded_file = st.file_uploader(
            "Lataa syötetiedosto", type="txt", label_visibility="collapsed"
        )
        if uploaded_file:
            syote_data = uploaded_file
    else:
        text_input = st.text_area(
            "Kirjoita tai liitä jäsenneltävä teksti tähän:", height=300
        )
        if text_input:
            syote_data = text_input

    with st.sidebar:
        st.header("Asetukset")
        top_k_valinta = st.number_input(
            "Haettavien jakeiden määrä per teema", 1, 100, 15
        )
        start_button = st.button(
            "Suorita haku", type="primary", disabled=(syote_data is None)
        )

    if start_button and syote_data:
        paaotsikko, sl, hakulauseet, otsikot = lue_syote_data(syote_data)
        if not hakulauseet:
            st.error("Syötteen jäsennys epäonnistui.")
        else:
            jae_kartta = defaultdict(dict)
            p_bar = st.progress(0, text="Aloitetaan...")
            total = len(hakulauseet)
            for i, (osio_nro, hakulause) in enumerate(hakulauseet.items()):
                haku = f"{paaotsikko}: {hakulause}"
                tulokset = etsi_merkityksen_mukaan(haku, top_k_valinta)
                jae_kartta[osio_nro]["jakeet"] = tulokset
                jae_kartta[osio_nro]["otsikko"] = otsikot.get(
                    osio_nro, hakulause
                )

                p_text = f"Käsitellään osiota {i+1}/{total}: {osio_nro}"
                p_bar.progress((i + 1) / total, text=p_text)

            st.session_state.final_report_md = luo_raportti_md(sl, jae_kartta)
            st.session_state.final_report_doc = luo_raportti_doc(sl, jae_kartta)
            st.session_state.processing_complete = True
            p_bar.empty()
            st.success("Haku suoritettu onnistuneesti!")

    if st.session_state.processing_complete:
        st.markdown(st.session_state.final_report_md)
        st.divider()
        st.subheader("Lataa raportti")
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "Lataa .txt-tiedostona",
                st.session_state.final_report_md,
                "raamattu_raportti.txt",
                "text/plain"
            )
        with col2:
            st.download_button(
                "Lataa .docx-tiedostona",
                st.session_state.final_report_doc,
                "raamattu_raportti.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


if __name__ == "__main__":
    main()